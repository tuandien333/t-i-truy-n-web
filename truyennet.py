import time
import os
import re
import sys
import shutil
import json
import urllib.request
import subprocess
import pickle
from selenium import webdriver
from selenium.webdriver.edge.options import Options as EdgeOptions
from selenium.webdriver.edge.service import Service as EdgeService
from webdriver_manager.microsoft import EdgeChromiumDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tqdm import tqdm
import glob

try:
    import win32com.client
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False

SECRET_KEY = 123 

class TruyenNetDownloader:
    def __init__(self):
        self.script_dir = os.path.dirname(os.path.abspath(__file__))
        self.root_dir = os.path.dirname(self.script_dir)
        
        self.temp_dir = os.path.join(self.root_dir, "temp")
        self.out_dir = os.path.join(self.root_dir, "Truyen_Tai_Ve")
        self.res_dir = os.path.join(self.root_dir, "Resources")
        self.pandoc_path = os.path.join(self.root_dir, "Pandoc", "pandoc.exe")
        self.history_file = os.path.join(self.root_dir, "truyennet_history.json")
        
        os.makedirs(self.temp_dir, exist_ok=True)
        os.makedirs(self.out_dir, exist_ok=True)
        os.makedirs(self.res_dir, exist_ok=True)
        
        self.driver = None
        self.book_title = "Truyen_Net"
        self.book_intro = ""
        self.cover_path = None
        self.chunk_chapters = [] 
        self.saved_parts = []    
        self.part_counter = 1    
        self.output_mode = '3'
        
        self.custom_font_name = "Times New Roman"
        self.check_custom_font()

    def check_custom_font(self):
        fonts = glob.glob(os.path.join(self.res_dir, "*.[to]tf"))
        if fonts:
            font_path = fonts[0]
            self.custom_font_name = os.path.splitext(os.path.basename(font_path))[0]

    def encrypt_data(self, data_obj):
        raw_bytes = pickle.dumps(data_obj)
        return bytearray([b ^ SECRET_KEY for b in raw_bytes])

    def decrypt_data(self, file_path):
        if not os.path.exists(file_path): return None
        with open(file_path, 'rb') as f: encrypted_bytes = f.read()
        raw_bytes = bytearray([b ^ SECRET_KEY for b in encrypted_bytes])
        try: return pickle.loads(raw_bytes)
        except: return None

    def load_history(self):
        if os.path.exists(self.history_file):
            try:
                with open(self.history_file, 'r', encoding='utf-8') as f: return json.load(f)
            except: return {}
        return {}

    def save_history(self, url, data):
        history = self.load_history()
        history[url] = data
        with open(self.history_file, 'w', encoding='utf-8') as f:
            json.dump(history, f, ensure_ascii=False, indent=4)

    def check_resume(self, url):
        history = self.load_history()
        if url in history:
            data = history[url]
            full_parts = data.get('parts', [])
            all_exist = True
            for part_file in full_parts:
                if not os.path.exists(os.path.join(self.temp_dir, part_file)):
                    all_exist = False; break
            if all_exist and data.get('last_chap_url'): return data
        return None

    def khoi_tao_driver(self):
        options = EdgeOptions()
        options.add_argument("--headless=new")
        options.add_argument("--disable-gpu")
        options.add_argument("--log-level=3")
        try:
            service = EdgeService(EdgeChromiumDriverManager().install())
            self.driver = webdriver.Edge(service=service, options=options)
        except:
            self.driver = webdriver.Edge(options=options)
        self.wait = WebDriverWait(self.driver, 20)

    def tai_anh_bia(self):
        try:
            img_elem = self.driver.find_element(By.CSS_SELECTOR, ".book img")
            url = img_elem.get_attribute("src")
            if url:
                save_path = os.path.join(self.temp_dir, "cover_tn.jpg")
                req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
                with urllib.request.urlopen(req) as resp, open(save_path, 'wb') as f:
                    f.write(resp.read())
                self.cover_path = save_path
        except: pass

    def lay_thong_tin_truyen(self):
        try:
            try: self.book_title = self.driver.find_element(By.CSS_SELECTOR, "h1.title").text.strip()
            except: pass
            self.tai_anh_bia()
            try:
                intro_elem = self.driver.find_element(By.CSS_SELECTOR, ".desc-text")
                self.book_intro = intro_elem.text.strip()
            except: pass
            
            try:
                btn = self.driver.find_element(By.CSS_SELECTOR, ".btn-primary.read-action")
                return btn.get_attribute("href")
            except: return None
        except: return None

    def loc_noi_dung(self):
        try:
            self.wait.until(EC.presence_of_element_located((By.ID, "chapter-c")))
            try:
                chap_title = self.driver.find_element(By.CSS_SELECTOR, ".chapter-title").text.strip()
            except: chap_title = "Chuong"

            self.driver.execute_script("document.querySelectorAll('.ads, div[id*=\"ads\"]').forEach(e => e.remove());")
            
            content_elem = self.driver.find_element(By.ID, "chapter-c")
            raw_text = content_elem.text
            lines = [l.strip() for l in raw_text.split('\n') if l.strip() and l.strip() != chap_title]
            
            return chap_title, "\n".join(lines)
        except: return None, None

    def luu_chunk_hien_tai(self):
        if not self.chunk_chapters: return
        safe_title = re.sub(r'[\\/*?:\"<>|]', '', self.book_title[:20]).strip()
        fname = f"tn_part_{self.part_counter}_{safe_title}_{int(time.time())}.bin"
        fpath = os.path.join(self.temp_dir, fname)

        data = {"chunk_data": self.chunk_chapters, "part_num": self.part_counter}
        with open(fpath, 'wb') as f: f.write(self.encrypt_data(data))
        
        self.saved_parts.append(fname)
        self.part_counter += 1
        
        log_data = {
            "title": self.book_title, "intro": self.book_intro,
            "cover": self.cover_path, "last_chap_url": self.driver.current_url,
            "parts": self.saved_parts
        }
        self.save_history(self.current_url_input, log_data)

    def apply_font(self, run):
        run.font.name = self.custom_font_name; run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.custom_font_name)

    def merge_va_xuat_file(self):
        if self.chunk_chapters: self.luu_chunk_hien_tai()
        if not self.saved_parts: return

        print("\nTruyện đang được hoàn thành, xin đợi giây lát....")
        master_doc = Document()
        section = master_doc.sections[0]
        section.page_width = Cm(21); section.page_height = Cm(29.7)
        section.top_margin = Cm(2); section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2)
        
        style = master_doc.styles['Normal']
        style.font.name = self.custom_font_name; style.font.size = Pt(13)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.custom_font_name)

        if self.cover_path and os.path.exists(self.cover_path):
            try:
                p = master_doc.add_paragraph(); p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.add_run().add_picture(self.cover_path, width=Cm(12))
            except: pass
        
        t = master_doc.add_heading(self.book_title, 0); t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in t.runs: self.apply_font(run)
        
        if self.book_intro:
            h = master_doc.add_heading("Gioi Thieu", 2)
            for run in h.runs: self.apply_font(run)
            p = master_doc.add_paragraph(self.book_intro); p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in p.runs: self.apply_font(run)
        master_doc.add_page_break()

        for part_file in self.saved_parts:
            data = self.decrypt_data(os.path.join(self.temp_dir, part_file))
            if not data: continue
            for item in data.get("chunk_data", []):
                h = master_doc.add_heading(item['title'], 1); h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in h.runs: self.apply_font(run)
                for line in item['content'].split('\n'):
                    p = master_doc.add_paragraph(line); p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    p.paragraph_format.first_line_indent = Cm(0.8)
                    for run in p.runs: self.apply_font(run)
                master_doc.add_page_break()

        safe_name = re.sub(r'[\\/*?:\"<>|]', '', self.book_title).strip()
        docx_path = os.path.join(self.out_dir, f"{safe_name}.docx")
        master_doc.save(docx_path)
        
        if self.output_mode in ['1', '3']:
            epub_path = os.path.join(self.out_dir, f"{safe_name}.epub")
            if os.path.exists(self.pandoc_path):
                try: subprocess.run([self.pandoc_path, docx_path, "-o", epub_path], creationflags=0x08000000); print(f"   [OK] EPUB: {epub_path}")
                except: pass
        
        if self.output_mode in ['2', '3'] and HAS_WIN32:
            pdf_path = os.path.join(self.out_dir, f"{safe_name}.pdf")
            try:
                w = win32com.client.Dispatch("Word.Application"); w.Visible = False
                d = w.Documents.Open(os.path.abspath(docx_path))
                d.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
                d.Close(); w.Quit(); print(f"   [OK] PDF: {pdf_path}")
            except: pass
            
        try: os.remove(docx_path)
        except: pass
        os.startfile(self.out_dir)

    def chay(self):
        url = sys.argv[1] if len(sys.argv) > 1 else ""
        if not url: url = input(">> Nhap link TruyenNet: ").strip()
        self.current_url_input = url

        print("\nChon dinh dang xuat ra:\n  (1) EPUB\n  (2) PDF\n  (3) EPUB + PDF")
        sel = input(">> Lua chon (1-3): ").strip()
        self.output_mode = sel if sel in ['1', '2', '3'] else '3'

        self.khoi_tao_driver()
        try:
            self.driver.get(url)
            time.sleep(2)
            
            history = self.check_resume(url)
            if history:
                print(f"\n[HISTORY] Phat hien truyen '{history['title']}'.")
                if input(">> Tiep tuc cap nhat? (y/n): ").lower() == 'y':
                    self.book_title = history['title']; self.book_intro = history.get('intro', '')
                    self.cover_path = history.get('cover'); self.saved_parts = history['parts']
                    self.part_counter = len(self.saved_parts) + 1
                    print("[*] Den chuong cuoi..."); self.driver.get(history['last_chap_url']); time.sleep(2)
                    try:
                        next_btns = self.driver.find_elements(By.CSS_SELECTOR, ".toolbox a.btn.blue")
                        if next_btns:
                             self.driver.execute_script("arguments[0].click();", next_btns[-1])
                        else:
                             print("[INFO] Het truyen."); self.merge_va_xuat_file(); return
                    except: return
                else: self.saved_parts = []
            
            if not self.saved_parts:
                if "/chuong-" not in url:
                    l = self.lay_thong_tin_truyen()
                    if l: self.driver.get(l)

            pbar = tqdm(unit=" ch", ncols=100)
            while True:
                t, c = self.loc_noi_dung()
                if t and c:
                    self.chunk_chapters.append({'title': t, 'content': c})
                    pbar.set_description(f"Tai: {t[:40]:<40}"); pbar.update(1)
                    if len(self.chunk_chapters) >= 100: self.luu_chunk_hien_tai(); self.chunk_chapters=[]
                
                try:
                    next_btns = self.driver.find_elements(By.CSS_SELECTOR, ".toolbox a.btn.blue")
                    if next_btns:
                        if "disabled" in next_btns[-1].get_attribute("class"): break
                        self.driver.execute_script("arguments[0].click();", next_btns[-1])
                    else: break
                except: break
        except Exception as e: print(f"\n[LOI] {e}")
        finally:
            pbar.close()
            if self.chunk_chapters: self.luu_chunk_hien_tai()
            self.merge_va_xuat_file()
            if self.driver: self.driver.quit()

if __name__ == "__main__":
    app = TruyenNetDownloader()
    app.chay()