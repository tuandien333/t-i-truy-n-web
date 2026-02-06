import time
import os
import re
import sys
import shutil
import json
import urllib.request
import subprocess
import pickle
from selenium import webdriver
from selenium.webdriver.edge.options import Options as EdgeOptions
from selenium.webdriver.edge.service import Service as EdgeService
from webdriver_manager.microsoft import EdgeChromiumDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tqdm import tqdm
import glob

try:
    import win32com.client
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False

SECRET_KEY = 123 

class TytNovelDownloader:
    def __init__(self):
        # --- CẤU HÌNH ĐƯỜNG DẪN ---
        self.script_dir = os.path.dirname(os.path.abspath(__file__))
        self.root_dir = os.path.dirname(self.script_dir)
        
        self.temp_dir = os.path.join(self.root_dir, "temp")
        self.out_dir = os.path.join(self.root_dir, "Truyen_Tai_Ve")
        self.res_dir = os.path.join(self.root_dir, "Resources")
        self.pandoc_path = os.path.join(self.root_dir, "Pandoc", "pandoc.exe")
        
        self.user_data_dir = os.path.join(self.root_dir, "UserData_Tyt")
        self.flag_file = os.path.join(self.root_dir, "tyt_saved.flag")
        self.history_file = os.path.join(self.root_dir, "tyt_history.json")
        
        os.makedirs(self.temp_dir, exist_ok=True)
        os.makedirs(self.out_dir, exist_ok=True)
        os.makedirs(self.res_dir, exist_ok=True)
        
        self.driver = None
        self.current_url = ""
        self.book_title = "Truyen_TytNovel"
        self.book_intro = ""
        self.cover_path = None
        self.chunk_chapters = [] 
        self.saved_parts = []    
        self.last_chap_url = ""  
        self.part_counter = 1    
        self.output_mode = '3' # Default both
        
        self.custom_font_name = "Times New Roman"
        self.check_custom_font()

    def check_custom_font(self):
        fonts = glob.glob(os.path.join(self.res_dir, "*.[to]tf"))
        if fonts:
            font_path = fonts[0]
            font_filename = os.path.basename(font_path)
            self.custom_font_name = os.path.splitext(font_filename)[0]
            # Da an thong bao FONT theo yeu cau

    def encrypt_data(self, data_obj):
        raw_bytes = pickle.dumps(data_obj)
        encrypted_bytes = bytearray([b ^ SECRET_KEY for b in raw_bytes])
        return encrypted_bytes

    def decrypt_data(self, file_path):
        if not os.path.exists(file_path): return None
        with open(file_path, 'rb') as f:
            encrypted_bytes = f.read()
        raw_bytes = bytearray([b ^ SECRET_KEY for b in encrypted_bytes])
        try:
            return pickle.loads(raw_bytes)
        except: return None

    def load_history(self):
        if os.path.exists(self.history_file):
            try:
                with open(self.history_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except: return {}
        return {}

    def save_history(self, url, data):
        history = self.load_history()
        history[url] = data
        with open(self.history_file, 'w', encoding='utf-8') as f:
            json.dump(history, f, ensure_ascii=False, indent=4)

    def check_resume(self, url):
        history = self.load_history()
        if url in history:
            data = history[url]
            full_parts = data.get('parts', [])
            all_exist = True
            for part_file in full_parts:
                if not os.path.exists(os.path.join(self.temp_dir, part_file)):
                    all_exist = False
                    break
            if all_exist and data.get('last_chap_url'):
                return data
        return None

    def khoi_tao_driver(self, hien_thi_trinh_duyet=False):
        # Da an thong bao SETUP theo yeu cau
        options = EdgeOptions()
        options.add_argument(f"--user-data-dir={self.user_data_dir}")
        options.add_argument("--profile-directory=Default")
        
        if not hien_thi_trinh_duyet:
            # Chế độ chạy ngầm (Headless)
            options.add_argument("--headless=new")
        else:
            # Chế độ hiện (để đăng nhập)
            options.add_argument("--start-maximized")

        options.add_argument("--disable-gpu")
        options.add_argument("--log-level=3")
        options.add_argument("--disable-blink-features=AutomationControlled")
        try:
            service = EdgeService(EdgeChromiumDriverManager().install())
            self.driver = webdriver.Edge(service=service, options=options)
        except:
            self.driver = webdriver.Edge(options=options)
        self.wait = WebDriverWait(self.driver, 20)

    def tai_anh_bia(self):
        try:
            img_elem = self.driver.find_element(By.CSS_SELECTOR, "img[itemprop='image']")
            url = img_elem.get_attribute("src")
            if url:
                ext = ".jpg"
                if ".png" in url.lower(): ext = ".png"
                elif ".webp" in url.lower(): ext = ".jpg"
                save_path = os.path.join(self.temp_dir, f"cover{ext}")
                req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
                with urllib.request.urlopen(req) as resp, open(save_path, 'wb') as f:
                    f.write(resp.read())
                self.cover_path = save_path
        except: pass

    def lay_thong_tin_truyen(self):
        try:
            try:
                self.book_title = self.driver.find_element(By.CSS_SELECTOR, "h5[itemprop='name']").text.strip()
                print(f"[*] Truyen: {self.book_title}")
            except: pass
            self.tai_anh_bia()
            try:
                intro_elem = self.driver.find_element(By.CSS_SELECTOR, "#scrollspyDesc .story-description")
                self.book_intro = intro_elem.text.replace('\n', ' ')
            except: pass
            try:
                start_btn = self.driver.find_element(By.CSS_SELECTOR, "a.btn-primary[title='Đọc Truyện']")
                return start_btn.get_attribute("href")
            except: return None
        except: return None

    def loc_noi_dung(self):
        try:
            self.wait.until(EC.presence_of_element_located((By.ID, "chapterContent")))
            try:
                title_elem = self.driver.find_element(By.CSS_SELECTOR, "div.text-center a.text-decoration-none[title*='Chương']")
                chap_title = title_elem.text.strip()
            except:
                try: chap_title = self.driver.find_element(By.CSS_SELECTOR, "h1, h2, h4").text.strip()
                except: chap_title = "Chuong khong ten"

            content_elem = self.driver.find_element(By.ID, "chapterContent")
            self.driver.execute_script("""
                var c = document.getElementById('chapterContent');
                if(c) { c.querySelectorAll('script, style, iframe, div[class*="ads"]').forEach(e => e.remove()); }
            """)
            raw_text = content_elem.text
            lines = [l.strip() for l in raw_text.split('\n') if l.strip()]
            return chap_title, "\n".join(lines)
        except: return None, None

    def luu_chunk_hien_tai(self):
        if not self.chunk_chapters: return
        safe_title = re.sub(r'[\\/*?:\"<>|]', '', self.book_title[:20]).strip()
        fname = f"part_{self.part_counter}_{safe_title}_{int(time.time())}.bin"
        fpath = os.path.join(self.temp_dir, fname)

        data_to_save = {
            "chunk_data": self.chunk_chapters,
            "part_num": self.part_counter
        }

        encrypted_content = self.encrypt_data(data_to_save)
        with open(fpath, 'wb') as f:
            f.write(encrypted_content)
        
        self.saved_parts.append(fname)
        self.part_counter += 1
        
        last_url = self.driver.current_url
        log_data = {
            "title": self.book_title,
            "intro": self.book_intro,
            "cover": self.cover_path,
            "last_chap_url": last_url,
            "parts": self.saved_parts
        }
        self.save_history(self.current_url, log_data)
        # Da an thong bao [SECURE] theo yeu cau

    def apply_font(self, run):
        run.font.name = self.custom_font_name
        run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.custom_font_name)

    def merge_va_xuat_file(self):
        if self.chunk_chapters:
            self.luu_chunk_hien_tai()
            self.chunk_chapters = []

        if not self.saved_parts:
            print("[INFO] Khong co du lieu.")
            return

        # Chi hien thong bao nay theo yeu cau
        print("\nTruyện đang được hoàn thành, xin đợi giây lát....")
        
        master_doc = Document()
        section = master_doc.sections[0]
        section.page_width = Cm(21); section.page_height = Cm(29.7)
        section.top_margin = Cm(2); section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2)
        
        style = master_doc.styles['Normal']
        style.font.name = self.custom_font_name
        style.font.size = Pt(13)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.custom_font_name)

        if self.cover_path and os.path.exists(self.cover_path):
            try:
                p = master_doc.add_paragraph(); p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.add_run().add_picture(self.cover_path, width=Cm(12))
            except: pass
        
        t_para = master_doc.add_heading(self.book_title, 0)
        t_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in t_para.runs: self.apply_font(run)
        
        if self.book_intro:
            h_intro = master_doc.add_heading("Gioi Thieu", 2)
            for run in h_intro.runs: self.apply_font(run)
            p = master_doc.add_paragraph(self.book_intro)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in p.runs: self.apply_font(run)

        master_doc.add_page_break()

        for part_file in self.saved_parts:
            part_path = os.path.join(self.temp_dir, part_file)
            data_obj = self.decrypt_data(part_path)
            if not data_obj: continue
            
            chapters = data_obj.get("chunk_data", [])
            for item in chapters:
                h = master_doc.add_heading(item['title'], 1)
                h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in h.runs: self.apply_font(run)

                for line in item['content'].split('\n'):
                    p = master_doc.add_paragraph(line)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    p.paragraph_format.line_spacing = 1.3
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.first_line_indent = Cm(0.8)
                    for run in p.runs: self.apply_font(run)
                master_doc.add_page_break()

        safe_name = re.sub(r'[\\/*?:\"<>|]', '', self.book_title).strip()
        docx_path = os.path.join(self.out_dir, f"{safe_name}.docx")
        master_doc.save(docx_path)
        
        # --- XU LY OUTPUT THEO LUA CHON ---
        
        # 1. EPUB
        if self.output_mode == '1' or self.output_mode == '3':
            epub_path = os.path.join(self.out_dir, f"{safe_name}.epub")
            if os.path.exists(self.pandoc_path):
                try: 
                    subprocess.run([self.pandoc_path, docx_path, "-o", epub_path], creationflags=0x08000000)
                    print(f"   [OK] EPUB: {epub_path}")
                except: pass

        # 2. PDF
        if self.output_mode == '2' or self.output_mode == '3':
            pdf_path = os.path.join(self.out_dir, f"{safe_name}.pdf")
            if HAS_WIN32:
                try:
                    w = win32com.client.Dispatch("Word.Application")
                    w.Visible = False
                    d = w.Documents.Open(os.path.abspath(docx_path))
                    d.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
                    d.Close()
                    w.Quit()
                    print(f"   [OK] PDF: {pdf_path}")
                except: pass

        # XOA FILE WORD CU THEO YEU CAU
        try:
            if os.path.exists(docx_path):
                os.remove(docx_path)
        except: pass

        os.startfile(self.out_dir)

    def xu_ly_cuoi_cung(self):
        print("\n" + "="*50)
        print("          QUAN LY TAI KHOAN")
        print("="*50)
        print("Ban co muon LUU LAI tai khoan nay?")
        ans = input(">> Lua chon (y/n): ").strip().lower()

        if ans == 'y':
            with open(self.flag_file, 'w') as f: f.write("saved")
            print("[OK] Da luu.")
        else:
            if os.path.exists(self.flag_file): os.remove(self.flag_file)
            if self.driver: self.driver.quit(); self.driver = None
            time.sleep(2)
            try:
                if os.path.exists(self.user_data_dir): shutil.rmtree(self.user_data_dir, ignore_errors=True)
                print("[OK] Da xoa.")
            except: pass

    def chay(self):
        url = sys.argv[1] if len(sys.argv) > 1 else ""
        need_login = False
        if len(sys.argv) > 2 and sys.argv[2] == "--login": need_login = True

        if not url: url = input(">> Nhap link TytNovel: ").strip()
        self.current_url = url

        # --- CHON DINH DANG ---
        print("\nChon dinh dang xuat ra:")
        print("  (1) EPUB")
        print("  (2) PDF")
        print("  (3) EPUB + PDF")
        sel = input(">> Lua chon (1-3): ").strip()
        if sel in ['1', '2', '3']: self.output_mode = sel
        else: self.output_mode = '3'

        # B1: KHOI DONG DRIVER
        # Neu can login -> Hien trinh duyet. Neu khong -> Chay ngam luon.
        self.khoi_tao_driver(hien_thi_trinh_duyet=need_login)
        
        try:
            if need_login:
                print("\n=== DANG NHAP THU CONG ===")
                self.driver.get("https://tytnovel.app/login")
                input(">> NHAN ENTER KHI DA DANG NHAP THANH CONG... ")
                
                # --- QUAN TRONG: CHUYEN SANG CHE DO AN (RESTART DRIVER) ---
                print("[INFO] Dang chuyen sang che do chay ngam...")
                self.driver.quit() # Tat trinh duyet hien
                with open(self.flag_file, 'w') as f: f.write("saved") # Luu co de nho da login
                
                # Khoi dong lai driver o che do ngam (Headless)
                self.khoi_tao_driver(hien_thi_trinh_duyet=False)

            self.driver.get(url)
            time.sleep(3)

            history_data = self.check_resume(url)
            
            if history_data:
                print(f"\n[HISTORY] Phat hien truyen '{history_data['title']}' da tung tai.")
                ans = input(">> Ban co muon TIEP TUC cap nhat khong? (y/n): ").lower()
                
                if ans == 'y':
                    self.book_title = history_data['title']
                    self.book_intro = history_data.get('intro', '')
                    self.cover_path = history_data.get('cover')
                    self.saved_parts = history_data['parts']
                    self.part_counter = len(self.saved_parts) + 1
                    
                    last_url = history_data['last_chap_url']
                    self.driver.get(last_url)
                    time.sleep(2)
                    try:
                        next_btn = self.driver.find_element(By.XPATH, "//a[contains(@class, 'btn') and .//i[contains(@class, 'bi-chevron-right')]]")
                        cls = next_btn.get_attribute("class")
                        if "disabled" in cls or "btn-outline-secondary" in cls:
                            print("[INFO] Chua co chuong moi.")
                            self.merge_va_xuat_file()
                            return
                        link_next = next_btn.get_attribute("href")
                        if link_next and "javascript" not in link_next: self.driver.get(link_next)
                        else:
                            self.driver.execute_script("arguments[0].click();", next_btn)
                            time.sleep(2)
                    except: return
                else:
                    self.saved_parts = []
            
            if not self.saved_parts:
                if "/chuong/" not in url:
                    link_start = self.lay_thong_tin_truyen()
                    if link_start: self.driver.get(link_start)
                    else: return

            pbar = tqdm(unit=" ch", ncols=100)
            while True:
                t, c = self.loc_noi_dung()
                if t and c:
                    self.chunk_chapters.append({'title': t, 'content': c})
                    pbar.set_description(f"Tai: {t[:40]:<40}")
                    pbar.update(1)
                    if len(self.chunk_chapters) >= 100:
                        self.luu_chunk_hien_tai()
                        self.chunk_chapters = [] 
                try:
                    next_btn = self.driver.find_element(By.XPATH, "//a[contains(@class, 'btn') and .//i[contains(@class, 'bi-chevron-right')]]")
                    if "disabled" in next_btn.get_attribute("class") or "btn-outline-secondary" in next_btn.get_attribute("class"):
                        break
                    link_next = next_btn.get_attribute("href")
                    if link_next and "javascript" not in link_next: self.driver.get(link_next)
                    else:
                        self.driver.execute_script("arguments[0].click();", next_btn)
                        time.sleep(1)
                except: break

        except Exception as e: print(f"\n[LOI CHINH] {e}")
        finally:
            pbar.close()
            if self.chunk_chapters: self.luu_chunk_hien_tai()
            self.merge_va_xuat_file()
            self.xu_ly_cuoi_cung()
            if self.driver: self.driver.quit()

if __name__ == "__main__":
    app = TytNovelDownloader()
    app.chay()