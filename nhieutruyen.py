import time
import re
import os
import sys
import subprocess
import urllib.request
import json
import random
from selenium import webdriver
from selenium.webdriver.edge.options import Options as EdgeOptions
from selenium.webdriver.chrome.options import Options as ChromeOptions
from selenium.webdriver.edge.service import Service as EdgeService
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from tqdm import tqdm

try:
    import win32com.client
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    from selenium_stealth import stealth
    HAS_STEALTH = True
except ImportError:
    HAS_STEALTH = False

# --- DANH SÁCH USER-AGENT GIẢ LẬP ---
USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:123.0) Gecko/20100101 Firefox/123.0"
]

class NhieuTruyenDownloader:
    def __init__(self):
        self.current_folder = os.path.dirname(os.path.abspath(__file__))
        self.temp_folder = os.path.join(self.current_folder, "temp")
        self.final_folder = os.path.join(self.current_folder, "truyện tải về")
        self.extension_folder = os.path.join(self.current_folder, "Extensions")
        self.pandoc_path = os.path.join(self.current_folder, "Pandoc", "pandoc.exe")
        self.font_path = os.path.join(self.current_folder, "Resources", "font.ttf")
        self.css_path = os.path.join(self.current_folder, "Resources", "style.css")
        self.raw_data_path = os.path.join(self.temp_folder, "full_data.json")

        # --- DANH SÁCH TỪ KHÓA RÁC (ĐÃ CẬP NHẬT MỚI) ---
        self.spam_keywords = [
            # --- MỚI BỔ SUNG ---
            "•´¯`•.  𝓉𝓇ộ𝓂 𝓉ừ 𝓃𝒽𝒾𝑒𝓊𝓉𝓇𝓊𝓎𝑒𝓃.𝒸𝓂  .•`¯´•",
            "𝓉𝓇ộ𝓂", 
            "𝓉ừ 𝓃𝒽𝒾𝑒𝓊𝓉𝓇𝓊𝓎𝑒𝓃.𝒸𝓂", 
            "đọ𝒸 ở 𝓃𝒽𝒾𝑒𝓊𝓉𝓇𝓊𝓎𝑒𝓃.𝒸𝓂",
            
            # Icon & Ký tự đặc biệt
            "🍑", "🎀", "🍪", "💙", "♡", "❤",
            
            # Các biến thể trước đây
            "đ͎ọ͎c͎ ͎t͎ạ͎i͎ ͎n͎h͎i͎e͎u͎t͎r͎u͎y͎e͎n͎.͎c͎o͎m͎",
            "đọ𝕔 𝕥ạ𝕚 𝕟𝕙𝕚𝕖𝕦𝕥𝕣𝕦𝕪𝕖𝕟.𝕔𝕠𝕞",
            "đ̾ọ̾c̾ ̾t̾ạ̾i̾ ̾n̾h̾i̾e̾u̾t̾r̾u̾y̾e̾n̾.̾c̾o̾m̾",
            "đ̾ọ̾c̾ ̾t̾ạ̾i̾", "̾n̾h̾i̾e̾u̾t̾r̾u̾y̾e̾n̾.̾c̾o̾m̾", 
            
            # Unicode Font lạ
            "•´¯`•. 𝓉𝓇ộ𝓂 .•`¯´•", "•´¯`•.  𝓉𝓇ộ𝓂   .•`¯´•", 
            "𝓉𝓇ộ𝓂 𝓉ừ 𝓃𝒽𝒾𝑒𝓊𝓉𝓇𝓊𝓎𝑒𝓃.𝒸🍪𝓂", 
            "𝓉ừ 𝓃𝒽𝒾𝑒𝓊𝓉𝓇𝓊𝓎𝑒𝓃.𝒸💙𝓂",
            "•´¯`•. 🎀 𝓉𝓇ộ𝓂", "🎀 .•`¯´•",
            "đọ𝒸 ở 𝓃𝒽𝒾𝑒𝓊𝓉𝓇𝓊𝓎𝑒𝓃.𝒸♡𝓂", "đọ𝒸 ở 𝓃𝒽𝒾𝑒𝓊𝓉𝓇𝓊𝓎𝑒𝓃.𝒸❤𝓂",
            "𝖙𝖗ộ𝖒", "𝖙ừ 𝖓𝖍𝖎𝖊𝖚𝖙𝖗𝖚𝖞𝖊𝖓.𝖈𝖔𝖒", "đọ𝚌 𝚝ạ𝚒", "𝚗𝚑𝚒𝚎𝚞𝚝𝚛𝚞𝚢𝚎𝚗.𝚌𝚘𝚖",
            "đọ𝔠 𝔱ạ𝔦", "𝔫𝔥𝔦𝔢𝔲𝔱𝔯𝔲𝔶𝔢𝔫.𝔠𝔬𝔪",
            "đọｃ ｔạｉ ｎｈｉｅｕｔｒｕｙｅｎ．ｃｏｍ", "ｔｒộｍ ｔừ ｎｈｉｅｕｔｒｕｙｅｎ．ｃｏｍ",
            
            # Text bị gạch chân/biến dạng
            "t̲r̲ộ̲m̲ ̲t̲ừ̲", "̲n̲h̲i̲e̲u̲t̲r̲u̲y̲e̲n̲.̲c̲o̲m̲",
            
            # Chữ in đậm/nghiêng
            "đọ𝙘 𝙩ạ𝙞", "𝙣𝙝𝙞𝙚𝙪𝙩𝙧𝙪𝙮𝙚𝙣.𝙘𝙤𝙢", "𝘵𝘳ộ𝘮", "𝘵ừ 𝘯𝘩𝘪𝘦𝘶𝘵𝘳𝘶𝘺𝘦𝘯.𝘤𝘰𝘮",
            
            # Tiếng Việt thường
            "trộm của", "trộm từ", "đọc tại", "đọc ở",
            "NhiềuTruyện.com(nhieutruyen.com)", "nhieutruyen.com", "NhiềuTruyện.com"
        ]

        if not os.path.exists(self.temp_folder): os.makedirs(self.temp_folder)
        if not os.path.exists(self.final_folder): os.makedirs(self.final_folder)
        if not os.path.exists(self.extension_folder): os.makedirs(self.extension_folder)
        
        self.driver = None; self.ten_truyen_goc = ""; self.cover_image_path = None

    def mo_trinh_duyet(self):
        if self.driver:
            try: self.driver.quit()
            except: pass
        
        random_ua = random.choice(USER_AGENTS)
        edge_opts = EdgeOptions(); chrome_opts = ChromeOptions()
        
        for opts in [edge_opts, chrome_opts]:
            opts.add_argument("--headless=new") 
            opts.add_argument("--log-level=3")
            opts.add_experimental_option('excludeSwitches', ['enable-logging'])
            opts.add_argument("--disable-blink-features=AutomationControlled")
            opts.add_argument(f"user-agent={random_ua}")
            
            if os.path.exists(self.extension_folder):
                for f in os.listdir(self.extension_folder):
                    if f.endswith(".crx"):
                        ext_path = os.path.abspath(os.path.join(self.extension_folder, f))
                        opts.add_extension(ext_path)

        old_stderr = sys.stderr
        try:
            sys.stderr = open(os.devnull, 'w')
            try:
                service = EdgeService(log_output=subprocess.DEVNULL)
                self.driver = webdriver.Edge(options=edge_opts, service=service)
            except:
                try:
                    service = ChromeService(log_output=subprocess.DEVNULL)
                    self.driver = webdriver.Chrome(options=chrome_opts, service=service)
                except: sys.stderr = old_stderr; print(f"[LỖI] Không tìm thấy Driver."); sys.exit()
        finally: sys.stderr = old_stderr
        
        if HAS_STEALTH:
            stealth(self.driver,
                languages=["en-US", "en"],
                vendor="Google Inc.",
                platform="Win32",
                webgl_vendor="Intel Inc.",
                renderer="Intel Iris OpenGL Engine",
                fix_hairline=True,
                )
        self.wait = WebDriverWait(self.driver, 20)

    def xoa_quang_cao(self):
        try: self.driver.execute_script("document.querySelectorAll('.ads, .adsbygoogle, iframe, div[id^=adm], .qc-row').forEach(e => e.remove());")
        except: pass

    # --- HÀM LỌC RÁC NÂNG CẤP ---
    def clean_spam_content(self, text):
        if not text: return ""
        
        for spam in self.spam_keywords:
            text = text.replace(spam, "")
            # Xóa cả khi nó bị xuống dòng cắt ngang
            text = text.replace(spam.replace(" ", "\n"), "") 
            text = text.replace(spam.replace(" ", "\n\n"), "")
            
        text = re.sub(r'\n\s*\n', '\n', text)
        return text.strip()

    def download_cover(self):
        print("[*] Đang quét ảnh bìa...")
        try:
            try: img_elem = self.wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "div.mb-3 img")))
            except: img_elem = self.driver.find_element(By.CSS_SELECTOR, "img.object-cover")
            if img_elem:
                img_url = img_elem.get_attribute("src")
                if img_url:
                    clean_name = re.sub(r'[\\/*?:\"<>|]', '', self.ten_truyen_goc).strip()
                    webp_path = os.path.join(self.temp_folder, f"{clean_name}.webp")
                    req = urllib.request.Request(img_url, headers={'User-Agent': 'Mozilla/5.0'})
                    with urllib.request.urlopen(req) as response:
                        with open(webp_path, 'wb') as f: f.write(response.read())
                    if HAS_PIL:
                        jpg_path = os.path.join(self.temp_folder, f"{clean_name}.jpg")
                        try:
                            im = Image.open(webp_path).convert("RGB")
                            im.save(jpg_path, "jpeg")
                            self.cover_image_path = os.path.abspath(jpg_path)
                            print(f"[OK] Đã tải ảnh: {os.path.basename(jpg_path)}")
                        except: self.cover_image_path = os.path.abspath(webp_path)
                    else: self.cover_image_path = os.path.abspath(webp_path)
        except: self.cover_image_path = None

    def luu_batch_word(self, batch_data, batch_index):
        if not batch_data: return
        doc = Document(); self.setup_docx(doc)
        doc.add_heading(f"PHẦN {batch_index}", 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for d in batch_data:
            doc.add_heading(d['title'], level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph("") 
            for line in d['content'].split('\n'):
                l = line.strip()
                if l: self.format_text(doc.add_paragraph(l))
            doc.add_page_break()
        filename = f"temp_part_{batch_index}.docx"
        save_path = os.path.join(self.temp_folder, filename)
        doc.save(save_path)
        print(f"\n[INFO] Đã lưu tạm: {filename}")

    def luu_data_vao_json(self, title, content):
        entry = {'title': title, 'content': content}
        with open(self.raw_data_path, 'a', encoding='utf-8') as f:
            f.write(json.dumps(entry, ensure_ascii=False) + "\n")

    def setup_docx(self, doc):
        s = doc.sections[0]; s.page_height = Cm(29.7); s.page_width = Cm(21.0)
        s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5); s.left_margin = Cm(3.0); s.right_margin = Cm(1.5)
    
    def format_text(self, p):
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY; p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.runs[0]; run.font.name = 'Times New Roman'; run.font.size = Pt(14)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    def tao_css_epub(self):
        with open(self.css_path, "w", encoding="utf-8") as f:
            f.write('@font-face { font-family: "MyFont"; src: url("font.ttf"); } body { font-family: "MyFont", serif; text-align: justify; } h1 { text-align: center; }')

    def convert_word_to_pdf_win32(self, input_docx, output_pdf):
        if not HAS_WIN32: return False
        abs_docx = os.path.abspath(input_docx); abs_pdf = os.path.abspath(output_pdf); word = None
        try:
            word = win32com.client.Dispatch("Word.Application"); word.Visible = False
            doc = word.Documents.Open(abs_docx); doc.SaveAs(abs_pdf, FileFormat=17); doc.Close()
            return True
        except: return False
        finally: 
            if word: word.Quit()

    def convert_epub(self, input_docx, output_epub):
        if not os.path.exists(self.pandoc_path): return False
        self.tao_css_epub()
        cmd = [self.pandoc_path, input_docx, "-o", output_epub, "--toc", "--metadata", f"title={self.ten_truyen_goc}", "--css", self.css_path]
        if os.path.exists(self.font_path): cmd.extend(["--epub-embed-font", self.font_path])
        if self.cover_image_path and os.path.exists(self.cover_image_path): cmd.extend(["--epub-cover-image", self.cover_image_path])
        try: 
            subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, creationflags=subprocess.CREATE_NO_WINDOW)
            return True
        except: return False

    def gop_va_xuat_file(self, mode):
        if not os.path.exists(self.raw_data_path): print("[WARN] Không có dữ liệu."); return
        print("\n" + "="*50); print(" ĐANG HỢP NHẤT FILE..."); print("="*50)
        doc = Document(); self.setup_docx(doc)
        if self.cover_image_path and os.path.exists(self.cover_image_path):
            try: p = doc.add_paragraph(); p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER; run = p.add_run(); run.add_picture(self.cover_image_path, width=Cm(14)); doc.add_page_break()
            except: pass
        doc.add_heading(self.ten_truyen_goc, 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER; doc.add_page_break()
        data_list = []
        with open(self.raw_data_path, 'r', encoding='utf-8') as f:
            for line in f:
                if line.strip(): data_list.append(json.loads(line))
        clean = self.ten_truyen_goc.lower().strip()
        
        spam_check = ["nhieutruyen", "đọc tại", "trộm từ"]
        
        for d in tqdm(data_list, desc="Gộp file", unit="ch", ncols=100):
            doc.add_heading(d['title'], level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER; doc.add_paragraph("")
            for line in d['content'].split('\n'):
                l = line.strip(); lo = l.lower()
                if not l or lo == clean: continue
                if any(x in lo for x in spam_check) and len(l) < 50: continue
                self.format_text(doc.add_paragraph(l))
            doc.add_page_break()
            
        filename = re.sub(r'[\\/*?:\"<>|]', '', self.ten_truyen_goc).strip()
        temp_docx = os.path.join(self.temp_folder, f"{filename}.docx")
        doc.save(temp_docx)
        final_epub = os.path.join(self.final_folder, f"{filename}.epub")
        final_pdf = os.path.join(self.final_folder, f"{filename}.pdf")
        if mode == '1': 
            print(f"[*] Đang xuất EPUB...");
            if self.convert_epub(temp_docx, final_epub): print(f"[THÀNH CÔNG] File lưu tại: {final_epub}")
        elif mode == '2': 
            print(f"[*] Đang xuất PDF...");
            if self.convert_word_to_pdf_win32(temp_docx, final_pdf): print(f"[THÀNH CÔNG] File lưu tại: {final_pdf}")
        try:
            for f in os.listdir(self.temp_folder):
                if f.startswith("temp_part_"): os.remove(os.path.join(self.temp_folder, f))
            if os.path.exists(temp_docx): os.remove(temp_docx)
            if self.cover_image_path and os.path.exists(self.cover_image_path): os.remove(self.cover_image_path)
            webp = self.cover_image_path.replace(".jpg", ".webp")
            if os.path.exists(webp): os.remove(webp)
            if os.path.exists(self.raw_data_path): os.remove(self.raw_data_path)
        except: pass

    def main(self, auto_url=None):
        if os.path.exists(self.raw_data_path): os.remove(self.raw_data_path)
        if auto_url:
            print(f"\n>> Đã nhận link từ Tìm kiếm: {auto_url}")
            url = auto_url
        else:
            url = input(">> Nhập link NhieuTruyen.com: ").strip()
        if not url: return

        print("\nCHỌN ĐỊNH DẠNG:\n1. EPUB (Điện thoại)\n2. PDF  (Máy tính)")
        choice = input(">> Nhập (1/2): ").strip()
        if choice not in ['1', '2']: choice = '1'

        limit_input = input(">> Tải bao nhiêu chương? (Nhập 0 để tải hết): ").strip()
        try: limit = int(limit_input)
        except: limit = 0

        self.mo_trinh_duyet(); self.driver.get(url)
        try: 
            try: self.ten_truyen_goc = self.driver.find_element(By.CSS_SELECTOR, "h1").text.strip()
            except: self.ten_truyen_goc = self.driver.title.split("|")[0].strip()
            print(f"[*] Truyện: {self.ten_truyen_goc}")
        except: print("[LỖI] Không tìm thấy tên truyện."); return

        self.download_cover()

        try:
            print("[*] Đang tìm nút 'Đọc Từ Đầu'...")
            read_btn = self.wait.until(EC.element_to_be_clickable((By.XPATH, "//a[contains(text(), 'Đọc Từ Đầu')]")))
            self.driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", read_btn); time.sleep(1); read_btn.click()
        except: print("[LỖI] Không tìm thấy nút Đọc Từ Đầu."); return
        
        pbar = tqdm(unit="ch", ncols=100)
        count = 0; batch_data = []; batch_count = 1
        
        try:
            while True:
                restart_threshold = random.randint(30, 45) 
                
                if count > 0 and count % restart_threshold == 0:
                    print(f"\n\n[*] Đã tải {count} chương. Đang lưu tạm...")
                    self.luu_batch_word(batch_data, batch_count)
                    
                    current_chapter_url = self.driver.current_url
                    batch_data = []; batch_count += 1
                    
                    print(f"[*] Đang 'Hạ nhiệt' (Xóa Cookie & Nghỉ 60s)...")
                    try: self.driver.delete_all_cookies() 
                    except: pass
                    self.driver.quit()
                    
                    for i in range(60, 0, -1):
                        sys.stdout.write(f"\r>> Nghỉ ngơi: {i} giây...   ")
                        sys.stdout.flush()
                        time.sleep(1)
                    print("\n[*] Đang kết nối lại...")
                    self.mo_trinh_duyet()
                    self.driver.get(current_chapter_url)

                try: 
                    t_elem = self.wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "h2.text-balance")))
                    t = t_elem.text.strip()
                except: t = "Chương Mới"
                
                c = ""
                try: 
                    c_elem = self.driver.find_element(By.ID, "chapter-content")
                    c_raw = c_elem.text
                    c = self.clean_spam_content(c_raw)
                except: c = "Lỗi nội dung."

                count += 1
                pbar.set_description(t[:30].ljust(30)); pbar.update(1)
                
                batch_data.append({'title': t, 'content': c})
                self.luu_data_vao_json(t, c)

                if limit > 0 and count >= limit:
                    print(f"\n[STOP] Đã tải đủ {limit} chương."); break
                
                try:
                    time.sleep(random.uniform(1.0, 3.0)) 
                    next_buttons = self.driver.find_elements(By.CSS_SELECTOR, "a.flex.items-center.justify-start.text-3xl")
                    if not next_buttons: break
                    nxt = next_buttons[-1]
                    self.driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", nxt)
                    self.driver.execute_script("arguments[0].click();", nxt)
                except: break
        except KeyboardInterrupt: pass
        except Exception as e: print(f"\n[LỖI] {e}")
        finally: pbar.close()
        
        if batch_data:
            self.luu_batch_word(batch_data, batch_count)

        self.gop_va_xuat_file(choice)
        if self.driver: self.driver.quit()

if __name__ == "__main__":
    app = NhieuTruyenDownloader()
    app.main()