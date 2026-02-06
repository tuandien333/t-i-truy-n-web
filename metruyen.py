import time
import os
import re
import sys
import shutil
import json
import urllib.request
import subprocess
import pickle
from selenium import webdriver
from selenium.webdriver.edge.options import Options as EdgeOptions
from selenium.webdriver.edge.service import Service as EdgeService
from webdriver_manager.microsoft import EdgeChromiumDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, NoSuchElementException
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tqdm import tqdm
import glob

try:
    import win32com.client
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False

# KEY MÃ HÓA (Đồng bộ với hệ thống)
SECRET_KEY = 123 

class MeTruyenDownloader:
    def __init__(self):
        # --- CẤU HÌNH ĐƯỜNG DẪN (ĐỒNG BỘ) ---
        # Script đang chạy trong thư mục Data
        self.script_dir = os.path.dirname(os.path.abspath(__file__))
        # Thư mục gốc (Cha của Data)
        self.root_dir = os.path.dirname(self.script_dir)
        
        # Các thư mục dữ liệu nằm ở Root
        self.temp_dir = os.path.join(self.root_dir, "temp")
        self.out_dir = os.path.join(self.root_dir, "Truyen_Tai_Ve")
        self.res_dir = os.path.join(self.root_dir, "Resources")
        self.ext_dir = os.path.join(self.root_dir, "Extensions")
        self.pandoc_path = os.path.join(self.root_dir, "Pandoc", "pandoc.exe")
        
        # File Log lịch sử
        self.history_file = os.path.join(self.root_dir, "metruyen_history.json")
        
        # Tạo thư mục nếu chưa có
        os.makedirs(self.temp_dir, exist_ok=True)
        os.makedirs(self.out_dir, exist_ok=True)
        os.makedirs(self.res_dir, exist_ok=True)
        
        self.driver = None
        self.current_url = ""
        self.book_title = "Truyen_MeTruyenCV"
        self.book_intro = ""
        self.cover_path = None
        self.chunk_chapters = [] 
        self.saved_parts = []    
        self.last_chap_url = ""  
        self.part_counter = 1    
        self.output_mode = '3'
        
        self.custom_font_name = "Times New Roman"
        self.check_custom_font()

    def check_custom_font(self):
        """Tìm font trong thư mục Resources của Root"""
        fonts = glob.glob(os.path.join(self.res_dir, "*.[to]tf"))
        if fonts:
            font_path = fonts[0]
            font_filename = os.path.basename(font_path)
            self.custom_font_name = os.path.splitext(font_filename)[0]
            # Da an thong bao ky thuat

    # --- MÃ HÓA / GIẢI MÃ ---
    def encrypt_data(self, data_obj):
        raw_bytes = pickle.dumps(data_obj)
        encrypted_bytes = bytearray([b ^ SECRET_KEY for b in raw_bytes])
        return encrypted_bytes

    def decrypt_data(self, file_path):
        if not os.path.exists(file_path): return None
        with open(file_path, 'rb') as f:
            encrypted_bytes = f.read()
        raw_bytes = bytearray([b ^ SECRET_KEY for b in encrypted_bytes])
        try:
            return pickle.loads(raw_bytes)
        except: return None

    # --- QUẢN LÝ LOG ---
    def load_history(self):
        if os.path.exists(self.history_file):
            try:
                with open(self.history_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except: return {}
        return {}

    def save_history(self, url, data):
        history = self.load_history()
        history[url] = data
        with open(self.history_file, 'w', encoding='utf-8') as f:
            json.dump(history, f, ensure_ascii=False, indent=4)

    def check_resume(self, url):
        history = self.load_history()
        if url in history:
            data = history[url]
            full_parts = data.get('parts', [])
            all_exist = True
            for part_file in full_parts:
                if not os.path.exists(os.path.join(self.temp_dir, part_file)):
                    all_exist = False
                    break
            if all_exist and data.get('last_chap_url'):
                return data
        return None

    # --- DRIVER ---
    def khoi_tao_driver(self):
        options = EdgeOptions()
        # Chạy ẩn để gọn gàng (Headless New)
        options.add_argument("--headless=new")
        options.add_argument("--disable-gpu")
        options.add_argument("--window-size=1920,1080")
        options.add_argument("--log-level=3")
        
        # Tự động nạp Extension (VD: uBlock) từ thư mục gốc/Extensions
        if os.path.exists(self.ext_dir):
            for ext_file in os.listdir(self.ext_dir):
                if ext_file.endswith(".crx"):
                    options.add_extension(os.path.join(self.ext_dir, ext_file))

        try:
            service = EdgeService(EdgeChromiumDriverManager().install())
            self.driver = webdriver.Edge(service=service, options=options)
        except:
            self.driver = webdriver.Edge(options=options)
        self.wait = WebDriverWait(self.driver, 20)

    # --- CRAWLER METRUYEN ---
    def tai_anh_bia(self):
        try:
            # Selector ảnh bìa MeTruyenCV
            img_elem = self.driver.find_element(By.CSS_SELECTOR, ".d-flex img, img.img-fluid")
            url = img_elem.get_attribute("src")
            if url:
                save_path = os.path.join(self.temp_dir, "cover_mt.jpg")
                req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
                with urllib.request.urlopen(req) as resp, open(save_path, 'wb') as f:
                    f.write(resp.read())
                self.cover_path = save_path
        except: pass

    def lay_thong_tin_truyen(self):
        try:
            # Lấy tên truyện
            try:
                self.book_title = self.driver.find_element(By.CSS_SELECTOR, "h1").text.strip()
            except: 
                self.book_title = "Truyen_MeTruyen"

            self.tai_anh_bia()
            
            # Lấy giới thiệu (nằm trong tab-content hoặc box giới thiệu)
            try:
                # Click nút "Xem thêm" nếu có
                try:
                    more_btn = self.driver.find_element(By.CLASS_NAME, "btn-show-more")
                    self.driver.execute_script("arguments[0].click();", more_btn)
                except: pass
                
                intro_elem = self.driver.find_element(By.CSS_SELECTOR, ".summary-content, .nh-read__description")
                self.book_intro = intro_elem.text.strip()
            except: pass

            # Lấy link đọc ngay
            try:
                # Nút "Đọc ngay" thường là button lớn
                read_btn = self.driver.find_element(By.XPATH, "//a[contains(text(), 'Đọc ngay') or contains(text(), 'Đọc từ đầu')]")
                return read_btn.get_attribute("href")
            except: 
                # Nếu không thấy nút đọc, thử tìm list chương
                try:
                    first_chap = self.driver.find_element(By.CSS_SELECTOR, "#nav-tab-content a")
                    return first_chap.get_attribute("href")
                except: return None
        except: return None

    def loc_noi_dung(self):
        """Lấy nội dung chương Metruyencv"""
        try:
            # Chờ load
            self.wait.until(EC.presence_of_element_located((By.ID, "article")))
            
            # Lấy tên chương (h2 hoặc .nh-read__title)
            try:
                chap_title = self.driver.find_element(By.CSS_SELECTOR, "h2.nh-read__title, h1").text.strip()
            except: chap_title = "Chuong khong ten"

            # Xóa rác trước khi lấy text
            self.driver.execute_script("""
                var c = document.getElementById('article');
                if(c) { 
                    c.querySelectorAll('iframe, script, style, div[style*="display:none"], .nh-read__alert').forEach(e => e.remove());
                }
            """)
            
            content_elem = self.driver.find_element(By.ID, "article")
            raw_text = content_elem.text
            
            # Lọc text rác
            lines = []
            for line in raw_text.split('\n'):
                l = line.strip()
                if l and "metruyencv" not in l.lower() and l != chap_title:
                    lines.append(l)
            
            return chap_title, "\n".join(lines)
        except: return None, None

    # --- LƯU FILE MÃ HÓA (BIN) ---
    def luu_chunk_hien_tai(self):
        if not self.chunk_chapters: return
        safe_title = re.sub(r'[\\/*?:\"<>|]', '', self.book_title[:20]).strip()
        fname = f"mt_part_{self.part_counter}_{safe_title}_{int(time.time())}.bin"
        fpath = os.path.join(self.temp_dir, fname)

        data_to_save = {
            "chunk_data": self.chunk_chapters,
            "part_num": self.part_counter
        }

        encrypted_content = self.encrypt_data(data_to_save)
        with open(fpath, 'wb') as f:
            f.write(encrypted_content)
        
        self.saved_parts.append(fname)
        self.part_counter += 1
        
        # Cập nhật log
        last_url = self.driver.current_url
        log_data = {
            "title": self.book_title,
            "intro": self.book_intro,
            "cover": self.cover_path,
            "last_chap_url": last_url,
            "parts": self.saved_parts
        }
        self.save_history(self.current_url, log_data)
        # Khong in ra man hinh [SECURE]

    # --- HÀM HỖ TRỢ FONT ---
    def apply_font(self, run):
        run.font.name = self.custom_font_name
        run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.custom_font_name)

    # --- GHÉP VÀ XUẤT FILE ---
    def merge_va_xuat_file(self):
        if self.chunk_chapters:
            self.luu_chunk_hien_tai()
            self.chunk_chapters = []

        if not self.saved_parts:
            print("[INFO] Khong co du lieu.")
            return

        print("\nTruyện đang được hoàn thành, xin đợi giây lát....")
        
        master_doc = Document()
        section = master_doc.sections[0]
        section.page_width = Cm(21); section.page_height = Cm(29.7)
        section.top_margin = Cm(2); section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2)
        
        style = master_doc.styles['Normal']
        style.font.name = self.custom_font_name
        style.font.size = Pt(13)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.custom_font_name)

        if self.cover_path and os.path.exists(self.cover_path):
            try:
                p = master_doc.add_paragraph(); p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.add_run().add_picture(self.cover_path, width=Cm(12))
            except: pass
        
        t_para = master_doc.add_heading(self.book_title, 0)
        t_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in t_para.runs: self.apply_font(run)
        
        if self.book_intro:
            h_intro = master_doc.add_heading("Gioi Thieu", 2)
            for run in h_intro.runs: self.apply_font(run)
            p = master_doc.add_paragraph(self.book_intro)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in p.runs: self.apply_font(run)

        master_doc.add_page_break()

        for part_file in self.saved_parts:
            part_path = os.path.join(self.temp_dir, part_file)
            data_obj = self.decrypt_data(part_path)
            if not data_obj: continue
            
            chapters = data_obj.get("chunk_data", [])
            for item in chapters:
                h = master_doc.add_heading(item['title'], 1)
                h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in h.runs: self.apply_font(run)

                for line in item['content'].split('\n'):
                    p = master_doc.add_paragraph(line)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    p.paragraph_format.line_spacing = 1.3
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.first_line_indent = Cm(0.8)
                    for run in p.runs: self.apply_font(run)
                master_doc.add_page_break()

        safe_name = re.sub(r'[\\/*?:\"<>|]', '', self.book_title).strip()
        docx_path = os.path.join(self.out_dir, f"{safe_name}.docx")
        master_doc.save(docx_path)
        
        # --- OUTPUT ---
        if self.output_mode == '1' or self.output_mode == '3': # EPUB
            epub_path = os.path.join(self.out_dir, f"{safe_name}.epub")
            if os.path.exists(self.pandoc_path):
                try: 
                    subprocess.run([self.pandoc_path, docx_path, "-o", epub_path], creationflags=0x08000000)
                    print(f"   [OK] EPUB: {epub_path}")
                except: pass

        if self.output_mode == '2' or self.output_mode == '3': # PDF
            pdf_path = os.path.join(self.out_dir, f"{safe_name}.pdf")
            if HAS_WIN32:
                try:
                    w = win32com.client.Dispatch("Word.Application")
                    w.Visible = False
                    d = w.Documents.Open(os.path.abspath(docx_path))
                    d.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
                    d.Close()
                    w.Quit()
                    print(f"   [OK] PDF: {pdf_path}")
                except: pass

        # Clean DOCX
        try:
            if os.path.exists(docx_path): os.remove(docx_path)
        except: pass

        os.startfile(self.out_dir)

    def chay(self):
        # Đọc tham số từ start.bat
        url = sys.argv[1] if len(sys.argv) > 1 else ""
        if not url: url = input(">> Nhap link MeTruyenCV: ").strip()
        self.current_url = url

        # Chọn chế độ
        print("\nChon dinh dang xuat ra:")
        print("  (1) EPUB")
        print("  (2) PDF")
        print("  (3) EPUB + PDF")
        sel = input(">> Lua chon (1-3): ").strip()
        if sel in ['1', '2', '3']: self.output_mode = sel
        else: self.output_mode = '3'

        self.khoi_tao_driver()
        
        try:
            self.driver.get(url)
            time.sleep(2)

            # Check Resume
            history_data = self.check_resume(url)
            if history_data:
                print(f"\n[HISTORY] Phat hien truyen cu '{history_data['title']}'")
                ans = input(">> Ban co muon TIEP TUC cap nhat khong? (y/n): ").lower()
                
                if ans == 'y':
                    self.book_title = history_data['title']
                    self.book_intro = history_data.get('intro', '')
                    self.cover_path = history_data.get('cover')
                    self.saved_parts = history_data['parts']
                    self.part_counter = len(self.saved_parts) + 1
                    
                    last_url = history_data['last_chap_url']
                    print(f"[*] Dang di den chuong cuoi...")
                    self.driver.get(last_url)
                    time.sleep(2)
                    
                    # Next sang chương mới
                    try:
                        # MeTruyen button GoNext
                        btns_next = self.driver.find_elements(By.CSS_SELECTOR, "button[data-x-bind='GoNext']")
                        if btns_next and btns_next[0].is_displayed():
                            self.driver.execute_script("arguments[0].click();", btns_next[0])
                            time.sleep(2)
                        else:
                            print("[INFO] Khong co chuong moi.")
                            self.merge_va_xuat_file()
                            return
                    except: return
                else:
                    self.saved_parts = []
            
            # Nếu tải mới
            if not self.saved_parts:
                if "/chuong-" not in url:
                    link_start = self.lay_thong_tin_truyen()
                    if link_start: self.driver.get(link_start)
                    else: 
                        print("[LOI] Khong tim thay link doc.")
                        return

            # Vòng lặp tải
            pbar = tqdm(unit=" ch", ncols=100)
            while True:
                try:
                    # Chờ text
                    self.wait.until(EC.presence_of_element_located((By.ID, "article")))
                    
                    t, c = self.loc_noi_dung()
                    if t and c:
                        self.chunk_chapters.append({'title': t, 'content': c})
                        pbar.set_description(f"Tai: {t[:40]:<40}")
                        pbar.update(1)

                        if len(self.chunk_chapters) >= 100:
                            self.luu_chunk_hien_tai()
                            self.chunk_chapters = []
                    
                    # Tìm nút Next (Metruyencv dùng button[data-x-bind='GoNext'])
                    btns_next = self.driver.find_elements(By.CSS_SELECTOR, "button[data-x-bind='GoNext']")
                    if btns_next and btns_next[0].is_displayed():
                        # Check disabled
                        if "disabled" in btns_next[0].get_attribute("class"):
                            break
                        self.driver.execute_script("arguments[0].click();", btns_next[0])
                        time.sleep(1) # Tránh bị block
                    else:
                        break
                        
                except TimeoutException:
                    print("\n[INFO] Timeout hoac Het truyen.")
                    break
                except Exception:
                    break

        except Exception as e: print(f"\n[LOI CHINH] {e}")
        finally:
            pbar.close()
            if self.chunk_chapters: self.luu_chunk_hien_tai()
            self.merge_va_xuat_file()
            if self.driver: self.driver.quit()

if __name__ == "__main__":
    app = MeTruyenDownloader()
    app.chay()